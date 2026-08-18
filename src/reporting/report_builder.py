"""
Generador de memoria de cálculo (Word) a partir de un resultado del
motor de cálculo IEEE 80.

Este módulo NO calcula nada -- solo formatea un DesignResult (y los
datos de entrada que lo generaron) en un documento profesional. Usa
python-docx (Python puro, sin dependencia de Node.js) para que el
resto del repo siga siendo instalable con un simple
`pip install -r requirements.txt`.

Consistente con el resto del proyecto: el documento generado incluye
siempre un espacio de firma del ingeniero responsable y, si el
DesignResult tiene notas (aproximaciones de suelo de dos capas, etc.),
esas notas se incluyen textualmente -- nunca se ocultan para que el
informe "se vea mejor".
"""

from dataclasses import dataclass
from datetime import date
from typing import Optional, Union, List, Dict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..engine.models import SoilModel, TwoLayerSoilModel, GridGeometry, FaultData, DesignResult

# Paleta simple y profesional
_COLOR_HEADER_BG = "1F4E5F"
_COLOR_PASS = "2E7D32"
_COLOR_FAIL = "C62828"


@dataclass
class ProjectInfo:
    project_name: str = "Proyecto sin nombre"
    location: str = ""
    client: str = ""
    prepared_by: str = ""
    report_date: Optional[date] = None

    def __post_init__(self):
        if self.report_date is None:
            self.report_date = date.today()


# ---------------------------------------------------------------------
# Helpers de formato (python-docx es bastante de bajo nivel para esto)
# ---------------------------------------------------------------------

def _shade_cell(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color: Optional[str] = None, size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_kv_table(doc: Document, rows: List[tuple], col_widths=(Cm(6), Cm(9))):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        row = table.add_row()
        _set_cell_text(row.cells[0], str(label), bold=True)
        _set_cell_text(row.cells[1], str(value))
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
    return table


def _add_results_table(doc: Document, rows: List[tuple]):
    """rows: lista de (nombre, valor, unidad)"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, text in enumerate(["Parámetro", "Valor", "Unidad"]):
        _set_cell_text(hdr[i], text, bold=True, color="FFFFFF")
        _shade_cell(hdr[i], _COLOR_HEADER_BG)
    for name, value, unit in rows:
        row = table.add_row().cells
        _set_cell_text(row[0], name)
        _set_cell_text(row[1], value)
        _set_cell_text(row[2], unit)
    return table


def _add_verdict_paragraph(doc: Document, label: str, passed: bool):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: {'CUMPLE' if passed else 'NO CUMPLE'}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(_COLOR_PASS if passed else _COLOR_FAIL)


# ---------------------------------------------------------------------
# Constructor principal
# ---------------------------------------------------------------------

def build_calculation_report(
    soil: Union[SoilModel, TwoLayerSoilModel],
    grid: GridGeometry,
    fault: FaultData,
    result: DesignResult,
    output_path: str,
    project_info: Optional[ProjectInfo] = None,
    visualization_reference: Optional[str] = None,
    static_view_paths: Optional[Dict[str, str]] = None,
) -> str:
    """
    Genera la memoria de cálculo en formato .docx.

    visualization_reference: ruta o nombre de archivo del dashboard 3D
      interactivo (generado por src/visual/plot3d.py), si existe, para
      mencionarlo como material complementario. No se embebe (es HTML
      interactivo, Word no puede mostrarlo).
    static_view_paths: dict {"potential": ruta.png, "touch": ruta.png,
      "step": ruta.png} generado por src/visual/static_plot3d.py --
      estas SÍ se embeben directamente en el documento como imágenes.
      Las claves ausentes simplemente no se incluyen.

    Devuelve la ruta del archivo generado.
    """
    project_info = project_info or ProjectInfo()

    doc = Document()

    # --- Estilos base ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # --- Portada / encabezado ---
    title = doc.add_heading("Memoria de Cálculo — Sistema de Puesta a Tierra", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Según IEEE Std 80 — Guide for Safety in AC Substation Grounding")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Proyecto", project_info.project_name),
        ("Ubicación", project_info.location or "—"),
        ("Cliente", project_info.client or "—"),
        ("Preparado por (asistido por sistema de agentes)", project_info.prepared_by or "—"),
        ("Fecha", project_info.report_date.strftime("%Y-%m-%d")),
    ])

    # --- Aviso obligatorio (human-in-the-loop) ---
    doc.add_paragraph()
    warning = doc.add_paragraph()
    warning_run = warning.add_run(
        "⚠️ Este documento fue generado con asistencia de un sistema de cálculo "
        "automatizado. Todo resultado debe ser revisado y firmado por un "
        "ingeniero eléctrico habilitado antes de su uso en un proyecto real. "
        "Ver espacio de firma al final de este documento."
    )
    warning_run.italic = True
    warning_run.font.size = Pt(9)
    warning_run.font.color.rgb = RGBColor.from_string(_COLOR_FAIL)

    # --- 1. Datos de entrada ---
    doc.add_heading("1. Datos de entrada", level=1)

    doc.add_heading("1.1 Modelo de suelo", level=2)
    if isinstance(soil, TwoLayerSoilModel):
        soil_rows = [
            ("Modelo", "Dos capas horizontales"),
            ("Resistividad capa superior (ρ1)", f"{soil.rho1:.1f} Ω·m"),
            ("Resistividad capa inferior (ρ2)", f"{soil.rho2:.1f} Ω·m"),
            ("Espesor capa superior (h1)", f"{soil.h1:.2f} m"),
            ("Factor de reflexión (K)", f"{soil.reflection_factor:.3f}"),
        ]
    else:
        soil_rows = [
            ("Modelo", "Uniforme"),
            ("Resistividad (ρ)", f"{soil.rho:.1f} Ω·m"),
        ]
    if soil.rho_s is not None:
        soil_rows.append(("Resistividad capa superficial (ρs)", f"{soil.rho_s:.1f} Ω·m"))
        soil_rows.append(("Espesor capa superficial (hs)", f"{soil.h_s:.2f} m"))
    _add_kv_table(doc, soil_rows)

    doc.add_heading("1.2 Geometría de la malla", level=2)
    _add_kv_table(doc, [
        ("Dimensiones (Lx × Ly)", f"{grid.Lx:.1f} m × {grid.Ly:.1f} m"),
        ("Profundidad de enterramiento (h)", f"{grid.h:.2f} m"),
        ("Diámetro del conductor (d)", f"{grid.d * 1000:.1f} mm"),
        ("Conductores paralelos (n_x × n_y)", f"{grid.n_x} × {grid.n_y}"),
        ("Número de varillas", str(grid.n_rods)),
        ("Longitud de cada varilla", f"{grid.L_rod:.2f} m" if grid.n_rods else "—"),
        ("Longitud total de conductor enterrado (Lt)", f"{grid.Lt:.1f} m"),
    ])

    doc.add_heading("1.3 Datos de falla", level=2)
    _add_kv_table(doc, [
        ("Corriente de falla simétrica (If)", f"{fault.If_sym:.0f} A"),
        ("Factor de división de corriente (Sf)", f"{fault.Sf:.2f}"),
        ("Duración de falla (tf)", f"{fault.tf:.2f} s"),
        ("Tiempo de exposición al choque (ts)", f"{fault.ts:.2f} s"),
        ("Relación X/R", f"{fault.X_R:.1f}"),
        ("Frecuencia", f"{fault.freq:.0f} Hz"),
        ("Factor de crecimiento futuro (Cp)", f"{fault.Cp:.2f}"),
    ])

    # --- 2. Resultados del cálculo ---
    doc.add_heading("2. Resultados del cálculo normativo", level=1)

    doc.add_heading("2.1 Corriente de diseño y resistencia", level=2)
    _add_results_table(doc, [
        ("Factor de decremento (Df)", f"{result.Df:.3f}", "—"),
        ("Corriente de malla (IG)", f"{result.IG:.1f}", "A"),
        ("Resistencia de malla (Rg)", f"{result.Rg:.3f}", "Ω"),
        ("Elevación de potencial de tierra (GPR)", f"{result.GPR:.1f}", "V"),
    ])

    doc.add_heading("2.2 Tensiones tolerables", level=2)
    _add_results_table(doc, [
        ("Factor de derating de superficie (Cs)", f"{result.Cs:.3f}", "—"),
        ("Tensión de contacto tolerable", f"{result.E_touch_tolerable:.1f}", "V"),
        ("Tensión de paso tolerable", f"{result.E_step_tolerable:.1f}", "V"),
    ])

    doc.add_heading("2.3 Factores geométricos", level=2)
    _add_results_table(doc, [
        ("Factor geométrico compuesto (n)", f"{result.n:.3f}", "—"),
        ("Factor de irregularidad (Ki)", f"{result.Ki:.3f}", "—"),
        ("Factor corrector Kii", f"{result.Kii:.3f}", "—"),
        ("Factor corrector por profundidad (Kh)", f"{result.Kh:.3f}", "—"),
        ("Factor geométrico de malla (Km)", f"{result.Km:.3f}", "—"),
        ("Factor geométrico de paso (Ks)", f"{result.Ks:.3f}", "—"),
        ("Longitud efectiva de malla (Lm)", f"{result.Lm:.1f}", "m"),
        ("Longitud efectiva de paso (Ls)", f"{result.Ls:.1f}", "m"),
    ])

    doc.add_heading("2.4 Tensiones resultantes y veredicto", level=2)
    _add_results_table(doc, [
        ("Tensión de malla (Em)", f"{result.Em:.1f}", "V"),
        ("Tensión de paso (Es)", f"{result.Es:.1f}", "V"),
    ])
    doc.add_paragraph()
    _add_verdict_paragraph(doc, "Tensión de malla (Em ≤ tensión de contacto tolerable)", result.mesh_ok)
    _add_verdict_paragraph(doc, "Tensión de paso (Es ≤ tensión de paso tolerable)", result.step_ok)
    doc.add_paragraph()
    _add_verdict_paragraph(doc, "VEREDICTO GENERAL DEL DISEÑO", result.passes)

    # --- 3. Notas y advertencias ---
    if result.notes:
        doc.add_heading("3. Notas y advertencias", level=1)
        for note in result.notes:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(note)

    # --- 4. Material complementario ---
    if visualization_reference or static_view_paths:
        doc.add_heading("4. Material complementario", level=1)
        doc.add_paragraph(
            "El siguiente material es exploratorio/visual -- complementa los "
            "valores normativos de la sección 2, que son los que rigen el "
            "veredicto. Se basa en un perfil de potencial calculado con una "
            "simplificación (corriente uniforme entre segmentos de la malla); "
            "ver notas de la sección 3 para más detalle sobre esa aproximación."
        )

        view_titles = {
            "potential": "Potencial de superficie",
            "touch": "Tensión de contacto aproximada",
            "step": "Tensión de paso aproximada",
        }
        if static_view_paths:
            for key in ("potential", "touch", "step"):
                path = static_view_paths.get(key)
                if path:
                    doc.add_heading(view_titles[key], level=2)
                    doc.add_picture(path, width=Cm(15))

        if visualization_reference:
            doc.add_paragraph(
                f"También hay disponible una versión interactiva (rotable, con "
                f"zoom) de estas mismas vistas: {visualization_reference}."
            )

    # --- 5. Firma ---
    doc.add_heading("Revisión y aprobación", level=1)
    doc.add_paragraph(
        "Este informe fue generado con asistencia de un sistema de cálculo "
        "automatizado (motor de cálculo determinístico IEEE 80 + agentes de "
        "verificación). Requiere revisión y firma de un ingeniero eléctrico "
        "habilitado antes de su uso en un proyecto real."
    )
    doc.add_paragraph()
    sig_table = _add_kv_table(doc, [
        ("Nombre del ingeniero responsable", "_______________________________"),
        ("Matrícula profesional", "_______________________________"),
        ("Firma", "_______________________________"),
        ("Fecha de revisión", "_______________________________"),
    ])

    doc.save(output_path)
    return output_path
