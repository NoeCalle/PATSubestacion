"""
Tests de la integración agentes <-> reportes (report_integration.py y
la generación automática desde coordinator.run_design_pipeline).

Usa el mismo patrón de FakeProvider que tests/test_agents.py (proveedor
LLM simulado, para no depender de una API key real), pero acá los
scripts SÍ incluyen tool calls reales -- lo que se está validando es
que el informe final se construye a partir del resultado que
efectivamente devolvió el motor de cálculo, nunca del texto del LLM.
"""

import json
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document

from src.agents.providers.base import LLMProvider, ProviderResponse, ToolCall
from src.agents.base_agent import Agent, AgentResult
from src.agents.session import DesignSession
from src.agents.tools import get_tool_dispatch
from src.agents.report_integration import extract_final_design_check, build_report_from_session
from src.agents.coordinator import run_design_pipeline
from src.reporting import ProjectInfo


# ---------- Helpers compartidos (mismo patrón que test_agents.py) ----------

class FakeProvider(LLMProvider):
    def __init__(self, script):
        self.script = list(script)
        self.calls = []

    def chat(self, system, messages, tools):
        self.calls.append({"system": system, "messages": messages, "tools": tools})
        if not self.script:
            raise AssertionError("FakeProvider se quedó sin respuestas scripteadas")
        return self.script.pop(0)

    def format_assistant_message(self, response):
        return {"role": "assistant", "content": response.text, "_tool_calls": response.tool_calls}

    def format_tool_result(self, tool_call_id, content):
        return {"role": "tool_result", "tool_call_id": tool_call_id, "content": content}


def text_response(text: str) -> ProviderResponse:
    return ProviderResponse(text=text, tool_calls=[])


def tool_call_response(name: str, input_: dict, call_id: str = "call_1") -> ProviderResponse:
    return ProviderResponse(text=None, tool_calls=[ToolCall(id=call_id, name=name, input=input_)])


VALID_UNIFORM_INPUT = {
    "rho": 100.0, "Lx": 60.0, "Ly": 60.0, "h": 0.5, "d": 0.01,
    "n_x": 7, "n_y": 7, "n_rods": 8, "L_rod": 3.0,
    "If_sym": 10000.0, "Sf": 0.6,
}


def make_real_tool_call(input_params: dict, tool_name: str = "run_design_check_uniform") -> dict:
    """Ejecuta la tool de verdad (motor de cálculo real) y arma la
    entrada tal como quedaría en AgentResult.tool_calls_made -- para
    que los tests de reportes trabajen sobre datos reales, no
    inventados."""
    func = get_tool_dispatch([tool_name])[tool_name]
    result_dict = func(**input_params)
    return {"tool": tool_name, "input": input_params, "result": json.dumps(result_dict, default=str)}


def full_document_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------- extract_final_design_check ----------

def test_extract_prefers_reviewer_call_over_designer_call():
    designer_call = make_real_tool_call({**VALID_UNIFORM_INPUT, "n_x": 5, "n_y": 5})
    reviewer_call = make_real_tool_call({**VALID_UNIFORM_INPUT, "n_x": 10, "n_y": 10})

    session = DesignSession(
        designer_results=[AgentResult(final_text="Propuesta", tool_calls_made=[designer_call])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=1,
    )

    extracted = extract_final_design_check(session)
    assert extracted is not None
    soil, grid, fault, result, source = extracted

    assert source == "reviewer"
    assert grid.n_x == 10 and grid.n_y == 10  # confirma que vino del reviewer_call, no del designer_call


def test_extract_falls_back_to_designer_when_reviewer_made_no_tool_call():
    designer_call = make_real_tool_call(VALID_UNIFORM_INPUT)

    session = DesignSession(
        designer_results=[AgentResult(final_text="Propuesta", tool_calls_made=[designer_call])],
        reviewer_results=[AgentResult(final_text="APROBADO (sin llamar la tool)", tool_calls_made=[])],
        iterations=1,
    )

    extracted = extract_final_design_check(session)
    assert extracted is not None
    _, _, _, _, source = extracted
    assert source == "designer"


def test_extract_returns_none_when_no_tool_calls_anywhere():
    session = DesignSession(
        designer_results=[AgentResult(final_text="Propuesta sin verificar", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[])],
        iterations=1,
    )

    assert extract_final_design_check(session) is None


def test_extract_returns_none_for_empty_session():
    assert extract_final_design_check(DesignSession()) is None


# ---------- build_report_from_session ----------

def test_build_report_from_session_uses_real_engine_values(tmp_path):
    reviewer_call = make_real_tool_call(VALID_UNIFORM_INPUT)
    expected_result = json.loads(reviewer_call["result"])

    session = DesignSession(
        designer_results=[AgentResult(final_text="Propuesta", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=2,
        final_verdict="APROBADO",
    )

    output_path = str(tmp_path / "report.docx")
    result_path = build_report_from_session(session, output_path)

    assert result_path == output_path
    assert os.path.exists(output_path)

    text = full_document_text(output_path)
    assert f"{expected_result['Em']:.1f}" in text
    assert f"{expected_result['Es']:.1f}" in text
    # Trazabilidad: debe decir de qué agente/iteración salió el resultado
    assert "reviewer" in text.lower()
    assert "iteración 2" in text.lower() or "iteracion 2" in text.lower()


def test_build_report_from_session_returns_none_without_tool_calls(tmp_path):
    session = DesignSession(
        designer_results=[AgentResult(final_text="Sin verificar", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[])],
        iterations=1,
    )

    output_path = str(tmp_path / "report.docx")
    result_path = build_report_from_session(session, output_path)

    assert result_path is None
    assert not os.path.exists(output_path)


def test_build_report_accepts_project_info(tmp_path):
    reviewer_call = make_real_tool_call(VALID_UNIFORM_INPUT)
    session = DesignSession(
        designer_results=[AgentResult(final_text="P", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=1,
    )

    info = ProjectInfo(project_name="Planta Piloto")
    output_path = str(tmp_path / "report.docx")
    build_report_from_session(session, output_path, project_info=info)

    text = full_document_text(output_path)
    assert "Planta Piloto" in text


# ---------- Vistas 3D automáticas ----------

def _count_embedded_images(docx_path: str) -> int:
    import zipfile
    with zipfile.ZipFile(docx_path) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


def test_build_report_embeds_3d_views_by_default(tmp_path):
    reviewer_call = make_real_tool_call(VALID_UNIFORM_INPUT)
    session = DesignSession(
        designer_results=[AgentResult(final_text="P", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=1,
    )

    output_path = str(tmp_path / "report.docx")
    build_report_from_session(session, output_path)  # include_3d_views=True por default

    assert _count_embedded_images(output_path) == 3
    text = full_document_text(output_path)
    assert "Material complementario" in text


def test_build_report_can_disable_3d_views(tmp_path):
    reviewer_call = make_real_tool_call(VALID_UNIFORM_INPUT)
    session = DesignSession(
        designer_results=[AgentResult(final_text="P", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=1,
    )

    output_path = str(tmp_path / "report.docx")
    build_report_from_session(session, output_path, include_3d_views=False)

    assert _count_embedded_images(output_path) == 0
    text = full_document_text(output_path)
    assert "Material complementario" not in text


def test_build_report_embeds_3d_views_for_two_layer_soil(tmp_path):
    """El suelo de dos capas debe convertirse a resistividad efectiva
    antes de calcular el perfil -- confirma que no crashea y que las
    vistas igual se generan."""
    reviewer_call = make_real_tool_call(
        {
            "rho1": 60.0, "rho2": 800.0, "h1": 4.0,
            "Lx": 60.0, "Ly": 60.0, "h": 0.5, "d": 0.01,
            "n_x": 7, "n_y": 7, "n_rods": 8, "L_rod": 3.0,
            "If_sym": 10000.0, "Sf": 0.6,
        },
        tool_name="run_design_check_two_layer",
    )
    session = DesignSession(
        designer_results=[AgentResult(final_text="P", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=1,
    )

    output_path = str(tmp_path / "report.docx")
    build_report_from_session(session, output_path)

    assert _count_embedded_images(output_path) == 3


def test_build_report_survives_3d_rendering_failure(tmp_path, monkeypatch):
    """Si el renderizado 3D falla, el informe se genera igual (sin esa
    sección) en vez de perderse todo el documento."""
    import src.agents.report_integration as report_integration_module

    def broken_render(*args, **kwargs):
        raise RuntimeError("fallo simulado de renderizado")

    monkeypatch.setattr(report_integration_module, "render_static_views", broken_render)

    reviewer_call = make_real_tool_call(VALID_UNIFORM_INPUT)
    session = DesignSession(
        designer_results=[AgentResult(final_text="P", tool_calls_made=[])],
        reviewer_results=[AgentResult(final_text="APROBADO", tool_calls_made=[reviewer_call])],
        iterations=1,
    )

    output_path = str(tmp_path / "report.docx")
    result_path = build_report_from_session(session, output_path)

    # El informe se generó igual, a pesar del fallo
    assert result_path == output_path
    assert os.path.exists(output_path)
    assert _count_embedded_images(output_path) == 0

    text = full_document_text(output_path)
    assert "No se pudieron generar las vistas 3D" in text
    assert "fallo simulado de renderizado" in text


# ---------- Integración completa: run_design_pipeline genera el .docx ----------

def test_pipeline_auto_generates_report_when_approved(tmp_path):
    provider = FakeProvider([
        text_response("Suelo uniforme, rho=100 Ohm*m."),                 # soil_agent
        text_response("Propuesta: malla 60x60m, 7x7, 8 varillas."),       # designer_agent (sin tool, simplificado)
        tool_call_response("run_design_check_uniform", VALID_UNIFORM_INPUT),  # reviewer_agent: llama la tool...
        text_response("APROBADO: verificado independientemente."),        # ...y aprueba
    ])

    output_path = str(tmp_path / "auto_report.docx")
    session = run_design_pipeline(
        provider,
        "Resistividad medida: 100 Ohm*m",
        "Falla 10kA, 60x60m disponibles",
        report_output_path=output_path,
        project_info=ProjectInfo(project_name="Subestación Auto-Reporte"),
    )

    assert session.final_verdict == "APROBADO"
    assert session.report_path == output_path
    assert os.path.exists(output_path)

    text = full_document_text(output_path)
    assert "Subestación Auto-Reporte" in text
    assert "CUMPLE" in text  # el veredicto normativo real queda en el documento


def test_pipeline_without_report_output_path_does_not_generate_report():
    provider = FakeProvider([
        text_response("Suelo uniforme."),
        text_response("Propuesta."),
        text_response("APROBADO."),
    ])

    session = run_design_pipeline(provider, "suelo", "proyecto")  # sin report_output_path

    assert session.report_path is None


def test_pipeline_report_path_is_none_when_nobody_called_the_tool(tmp_path):
    """Si ni el diseñador ni el revisor llegaron a llamar la tool (ej.
    el LLM no siguió las instrucciones), el pipeline no debe generar
    un documento con datos inventados -- report_path debe quedar None."""
    provider = FakeProvider([
        text_response("Suelo uniforme."),
        text_response("Propuesta sin verificar con la herramienta."),
        text_response("APROBADO (sin evidencia real)."),
    ])

    output_path = str(tmp_path / "report.docx")
    session = run_design_pipeline(
        provider, "suelo", "proyecto", report_output_path=output_path
    )

    assert session.report_path is None
    assert not os.path.exists(output_path)
