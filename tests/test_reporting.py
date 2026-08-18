"""
Tests del módulo de reportes.

Verifica estructura y contenido del .docx generado (usando python-docx
para leerlo de vuelta), no el renderizado visual en sí -- eso se
verificó manualmente renderizando a PDF durante el desarrollo (ver
notas de la conversación / proceso de construcción).
"""

import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document

from src.engine.models import SoilModel, TwoLayerSoilModel, GridGeometry, FaultData
from src.engine.design_check import run_design_check, run_design_check_two_layer
from src.reporting.report_builder import build_calculation_report, ProjectInfo


def sample_uniform_case():
    soil = SoilModel(rho=100.0)
    grid = GridGeometry(
        Lx=60.0, Ly=60.0, h=0.5, d=0.01, n_x=7, n_y=7,
        n_rods=8, L_rod=3.0, rods_on_perimeter=True,
    )
    fault = FaultData(If_sym=10000.0, Sf=0.6, tf=0.5, ts=0.5, X_R=15.0, freq=60.0)
    result = run_design_check(soil, grid, fault, body_kg=50)
    return soil, grid, fault, result


def full_document_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_report_generates_file(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")

    returned_path = build_calculation_report(soil, grid, fault, result, output_path)

    assert returned_path == output_path
    assert os.path.exists(output_path)
    assert os.path.getsize(output_path) > 1000


def test_report_contains_key_result_values(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")
    build_calculation_report(soil, grid, fault, result, output_path)

    text = full_document_text(output_path)

    assert f"{result.Em:.1f}" in text
    assert f"{result.Es:.1f}" in text
    assert f"{result.GPR:.1f}" in text
    assert f"{result.Rg:.3f}" in text


def test_report_shows_correct_verdict_text(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")
    build_calculation_report(soil, grid, fault, result, output_path)

    text = full_document_text(output_path)

    # Sabemos por otros tests/ejemplos que este caso NO cumple
    assert not result.passes
    assert "NO CUMPLE" in text
    assert "CUMPLE" in text  # aparece dentro de "NO CUMPLE" también, solo confirma presencia


def test_report_includes_all_notes_verbatim(tmp_path):
    """Las notas de advertencia (ej. aproximaciones) nunca deben omitirse
    del informe, incluso si son 'negativas' para el diseño."""
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")
    build_calculation_report(soil, grid, fault, result, output_path)

    text = full_document_text(output_path)

    assert len(result.notes) > 0  # este caso genera notas (no cumple)
    for note in result.notes:
        assert note in text


def test_report_includes_signature_section(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")
    build_calculation_report(soil, grid, fault, result, output_path)

    text = full_document_text(output_path)

    assert "ingeniero" in text.lower()
    assert "Matrícula" in text or "matrícula" in text.lower()


def test_report_two_layer_soil_shows_reflection_factor(tmp_path):
    soil = TwoLayerSoilModel(rho1=60.0, rho2=800.0, h1=4.0)
    grid = GridGeometry(
        Lx=60.0, Ly=60.0, h=0.5, d=0.01, n_x=7, n_y=7,
        n_rods=8, L_rod=3.0, rods_on_perimeter=True,
    )
    fault = FaultData(If_sym=10000.0, Sf=0.6, tf=0.5, ts=0.5, X_R=15.0, freq=60.0)
    result = run_design_check_two_layer(soil, grid, fault, body_kg=50)

    output_path = str(tmp_path / "report_two_layer.docx")
    build_calculation_report(soil, grid, fault, result, output_path)

    text = full_document_text(output_path)

    assert "Dos capas" in text
    assert f"{soil.reflection_factor:.3f}" in text
    # La nota de aproximación de resistividad efectiva debe estar presente
    assert any("aproximación" in n.lower() for n in result.notes)
    assert "aproximación" in text.lower()


def test_report_uses_custom_project_info(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")

    info = ProjectInfo(
        project_name="Subestación Norte 220kV",
        location="Arequipa, Perú",
        client="Empresa Eléctrica de Prueba",
        prepared_by="Sistema de agentes IEEE 80",
    )
    build_calculation_report(soil, grid, fault, result, output_path, project_info=info)

    text = full_document_text(output_path)
    assert "Subestación Norte 220kV" in text
    assert "Arequipa, Perú" in text
    assert "Empresa Eléctrica de Prueba" in text


def test_report_mentions_visualization_reference_when_provided(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")

    build_calculation_report(
        soil, grid, fault, result, output_path,
        visualization_reference="potential_dashboard.html",
    )

    text = full_document_text(output_path)
    assert "potential_dashboard.html" in text


def test_report_omits_visualization_section_when_not_provided(tmp_path):
    soil, grid, fault, result = sample_uniform_case()
    output_path = str(tmp_path / "report.docx")

    build_calculation_report(soil, grid, fault, result, output_path)

    text = full_document_text(output_path)
    assert "Material complementario" not in text


def _count_embedded_images(docx_path: str) -> int:
    """Cuenta las imágenes efectivamente embebidas en el .docx (no solo
    mencionadas en texto) -- un .docx es un .zip, las imágenes viven
    en word/media/."""
    import zipfile
    with zipfile.ZipFile(docx_path) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


def test_report_embeds_static_view_images(tmp_path):
    from src.visual.static_plot3d import render_static_views
    from src.engine.potential_profile import (
        compute_surface_potential, touch_voltage_field, step_voltage_field,
    )

    soil, grid, fault, result = sample_uniform_case()

    field = compute_surface_potential(soil, grid, IG=result.IG, margin=5.0, sample_resolution=4.0)
    touch_field = touch_voltage_field(field.V, gpr_reference=result.GPR)
    step_field = step_voltage_field(field.X, field.Y, field.V)

    static_paths = render_static_views(
        field, grid, output_dir=str(tmp_path / "views"),
        touch_field=touch_field, step_field=step_field,
    )

    output_path = str(tmp_path / "report_with_images.docx")
    build_calculation_report(
        soil, grid, fault, result, output_path,
        static_view_paths=static_paths,
    )

    text = full_document_text(output_path)
    assert "Material complementario" in text
    assert "Potencial de superficie" in text
    assert "Tensión de contacto aproximada" in text
    assert "Tensión de paso aproximada" in text

    # Lo importante: las imágenes deben estar REALMENTE embebidas, no
    # solo mencionadas en el texto.
    assert _count_embedded_images(output_path) == 3


def test_report_embeds_only_provided_views(tmp_path):
    """Si solo se pasa la vista de potencial (sin touch/step), el
    informe no debe mencionar ni intentar embeber las otras."""
    from src.visual.static_plot3d import render_static_views
    from src.engine.potential_profile import compute_surface_potential

    soil, grid, fault, result = sample_uniform_case()
    field = compute_surface_potential(soil, grid, IG=result.IG, margin=5.0, sample_resolution=4.0)

    static_paths = render_static_views(field, grid, output_dir=str(tmp_path / "views"))

    output_path = str(tmp_path / "report_potential_only.docx")
    build_calculation_report(soil, grid, fault, result, output_path, static_view_paths=static_paths)

    assert _count_embedded_images(output_path) == 1
    text = full_document_text(output_path)
    assert "Tensión de contacto aproximada" not in text


def test_report_mentions_both_static_images_and_interactive_reference(tmp_path):
    from src.visual.static_plot3d import render_static_views
    from src.engine.potential_profile import compute_surface_potential

    soil, grid, fault, result = sample_uniform_case()
    field = compute_surface_potential(soil, grid, IG=result.IG, margin=5.0, sample_resolution=4.0)
    static_paths = render_static_views(field, grid, output_dir=str(tmp_path / "views"))

    output_path = str(tmp_path / "report_both.docx")
    build_calculation_report(
        soil, grid, fault, result, output_path,
        static_view_paths=static_paths,
        visualization_reference="potential_dashboard.html",
    )

    assert _count_embedded_images(output_path) == 1
    text = full_document_text(output_path)
    assert "potential_dashboard.html" in text
